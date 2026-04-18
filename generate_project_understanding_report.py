from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"D:\proctoring AI")
BACKEND_ROOT = ROOT / "Proctoring-AI-BE-M4" / "Proctoring-AI-BE-M4"
STUDENT_ROOT = ROOT / "Proctoring-AI-FE-M4" / "Proctoring-AI-FE-M4"
ADMIN_ROOT = ROOT / "Proctoring-AI-Admin"
OUTPUT_PATH = ROOT / "PROJECT_UNDERSTANDING_REPORT.docx"

SOURCE_EXTENSIONS = {
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".css",
    ".json",
    ".md",
    ".yml",
    ".yaml",
}
SKIP_PARTS = {
    ".git",
    "__pycache__",
    "node_modules",
    "dist",
    "build",
    ".vite",
    ".idea",
    ".vscode",
    ".pytest_cache",
    ".mypy_cache",
    ".venv",
    "venv",
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def list_source_files(root: Path) -> list[Path]:
    files: list[Path] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if any(part in SKIP_PARTS for part in path.parts):
            continue
        if path.suffix.lower() not in SOURCE_EXTENSIONS:
            continue
        files.append(path)
    return files


def line_count(path: Path) -> int:
    try:
        with path.open("r", encoding="utf-8", errors="ignore") as handle:
            return sum(1 for _ in handle)
    except OSError:
        return 0


def load_json(path: Path) -> dict:
    try:
        return json.loads(read_text(path))
    except Exception:
        return {}


def extract_backend_endpoints(router_path: Path) -> list[dict[str, str]]:
    text = read_text(router_path)
    lines = text.splitlines()
    endpoints: list[dict[str, str]] = []
    route_re = re.compile(r'@router\.(get|post|patch|put|delete)\(\s*["\']([^"\']+)["\']')
    func_re = re.compile(r"^\s*(?:async\s+def|def)\s+([A-Za-z_][A-Za-z0-9_]*)")
    for index, line in enumerate(lines):
        match = route_re.search(line)
        if not match:
            continue
        method = match.group(1).upper()
        route_path = match.group(2)
        handler = ""
        for candidate in lines[index + 1 : index + 12]:
            func_match = func_re.search(candidate)
            if func_match:
                handler = func_match.group(1)
                break
        endpoints.append(
            {
                "router_file": router_path.name,
                "method": method,
                "path": route_path,
                "handler": handler or "unknown_handler",
            }
        )
    return endpoints


def extract_python_classes(path: Path) -> list[str]:
    text = read_text(path)
    return re.findall(r"^class\s+([A-Za-z_][A-Za-z0-9_]*)", text, flags=re.MULTILINE)


def extract_react_routes(app_file: Path) -> list[dict[str, str]]:
    text = read_text(app_file)
    route_re = re.compile(
        r'<Route\s+path="([^"]+)"\s+element=\{<([A-Za-z0-9_]+)(?:\s*/)?>\}\s*/?>'
    )
    routes: list[dict[str, str]] = []
    for match in route_re.finditer(text):
        routes.append({"path": match.group(1), "component": match.group(2)})
    return routes


def extract_export_names(path: Path) -> list[str]:
    text = read_text(path)
    names = set()
    for pattern in [
        r"export\s+default\s+function\s+([A-Za-z_][A-Za-z0-9_]*)",
        r"export\s+function\s+([A-Za-z_][A-Za-z0-9_]*)",
        r"export\s+default\s+([A-Za-z_][A-Za-z0-9_]*)",
        r"const\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*\(",
        r"function\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(",
    ]:
        for match in re.findall(pattern, text):
            names.add(match)
    return sorted(names)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        add_paragraph(doc, item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Project Understanding Report\nAI Proctoring Platform")
    run.bold = True
    run.font.size = Pt(22)

    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer.add_run(
        "Backend + Student Frontend + Admin Frontend\nGenerated from the current repository state"
    ).font.size = Pt(12)

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(
        "Workspace: D:\\proctoring AI\nDate: 2026-04-18\nAudience: Lecturer / Evaluator / Reviewer"
    ).font.size = Pt(11)

    doc.add_page_break()


def app_kind_summary() -> list[str]:
    return [
        "This repository is not a single web app. It is a coordinated exam ecosystem composed of a FastAPI backend, a student-facing React application, and an admin-facing React application used for live monitoring, results analysis, and exam configuration.",
        "At a practical level, the backend is the operational center of the system. It stores users and exam sessions, issues JWTs, hosts REST endpoints, manages WebSocket connections, runs the computer-vision proctoring pipeline, records violations, generates evidence artifacts, and exposes administrative summary data.",
        "The student frontend behaves like an exam client. It handles authentication, guided face enrollment, identity verification, exam lobby flows, browser and device checks, screen sharing, the active exam page, violation feedback, timing, submission, and final summary presentation.",
        "The admin frontend behaves like an operations console. It exposes a dashboard for live rooms, a results and evidence review workflow, a student summary view, settings management, and exam creation utilities. This makes the project closer to a small product suite than a classroom CRUD application.",
        "The architecture shows two important design goals. First, the team aimed for real-time proctoring rather than post-exam batch analysis. Second, the project tries to combine strong identity verification with exam-session governance, which is why login, liveness, face comparison, session rules, and evidence storage are tightly connected.",
        "A reviewer should therefore read the codebase in terms of flows rather than isolated files: onboarding flow, identity flow, exam start flow, frame-processing flow, policy-enforcement flow, live-monitoring flow, evidence-review flow, and exam-result flow.",
    ]


def architecture_summary() -> list[str]:
    return [
        "The runtime topology is straightforward but ambitious. Docker Compose provisions PostgreSQL, Redis, MinIO-compatible object storage, the FastAPI backend, and Adminer. The two frontends run independently through Vite development servers on separate ports. This split makes each application easier to evolve while still keeping the local developer experience understandable.",
        "The backend startup path in main.py initializes database state, storage infrastructure, middleware, API routers, and WebSocket endpoints. That means the backend is not simply exposing models through REST; it is orchestrating both synchronous request-response traffic and long-lived real-time processing channels.",
        "The WebSocket layer is especially important to understanding system behavior. Student clients send image frames to the backend during the active exam. The backend passes frames through a detector pipeline, emits logs and processing acknowledgements, and persists major events to the database and storage systems. The admin side consumes summarized live data rather than raw camera transport.",
        "The backend uses a service-heavy design. Routers define public contracts, but operational rules sit in services such as detection_service, storage_service, warmup_service, grading_service, log_service, and termination_policy_service. This is a healthy pattern because it keeps controllers from becoming the only place where business logic exists.",
        "Configuration is split between docker-compose, environment variables, settings.py, and detection_config.py. This is meaningful in a proctoring platform because deployment, computer-vision thresholds, and policy tuning all materially change runtime behavior. The report therefore treats configuration as part of the architecture rather than as an implementation footnote.",
        "A useful mental model is to think of the project as four cooperating layers: interface layer, orchestration layer, policy layer, and evidence layer. The frontends make up the interface layer; FastAPI routers and WebSocket handlers make up the orchestration layer; detector logic, identity checks, and termination policy make up the policy layer; database rows and MinIO evidence files make up the evidence layer.",
        "From an academic evaluation perspective, the project demonstrates system integration across UI, backend APIs, concurrency, media capture, machine-learning inference, persistent audit trails, and role-based administration. The engineering challenge is not only writing components but keeping all of those moving parts aligned in one exam lifecycle.",
    ]


def backend_deep_dive() -> list[tuple[str, list[str]]]:
    return [
        (
            "Backend Role In The System",
            [
                "The backend is the authoritative engine of the project. It is responsible for identity, session state, policy enforcement, event logging, evidence persistence, results aggregation, and machine-vision orchestration. In many academic projects the backend is only a data API, but in this system it is where the product’s most distinctive behavior lives.",
                "The router inventory confirms that the backend supports more than candidate exam attempts. It also provides admin-oriented endpoints, settings management, LTI entry points, observability contracts, evidence retrieval, and exam authoring. This means the backend functions as a platform service shared by two frontends and likely by LMS integrations.",
                "Because the backend owns both the exam record and the proctoring decisions, it becomes the single source of truth when the frontends disagree or disconnect. That is why terminated sessions remain terminated even if the browser reconnects later. This choice is operationally correct for auditability, though it can create friction during testing unless reset utilities exist.",
            ],
        ),
        (
            "Startup, Configuration, And Environment",
            [
                "The compose file shows a modern containerized stack: PostgreSQL for transactional data, Redis for cache or session-style support, MinIO for object storage, Adminer for inspection, and a Gunicorn-hosted FastAPI process for the application itself. This layout is appropriate for a system that stores both structured exam metadata and heavier evidence assets.",
                "settings.py centralizes operational variables such as JWT behavior, frontend URLs, WebSocket base URL, storage credentials, Redis location, LTI fields, and enforcement thresholds. detection_config.py separately centralizes detector-specific sensitivity values. That split is a good sign because product-level configuration and model-level tuning usually evolve at different speeds.",
                "One important observation for a lecturer is that repository documentation has drifted over time. The backend README still references older assumptions such as MySQL and older setup guidance, while the current runtime stack in Docker Compose clearly favors PostgreSQL and container orchestration. This does not invalidate the project, but it is worth noting as a maintenance lesson.",
                "The Dockerfile also tells a story about real-world deployment pain. It preinstalls heavy CPU-only ML dependencies, downloads DeepFace weights during build time, and defaults to conservative Gunicorn concurrency. Those decisions reflect experience with worker cold starts, timeouts, and large-model initialization issues.",
            ],
        ),
        (
            "HTTP API Structure",
            [
                "The backend routers are organized around capability domains. auth.py handles account operations, identity enrollment, and authentication variants. exam.py is the largest operational surface and covers session lifecycle, logging, summary, evidence, exam creation, room monitoring, and admin reporting. settings.py exposes health and configuration endpoints. observability.py surfaces policy information. lti.py covers LMS-related integration points.",
                "This organization is pragmatic rather than academically pure REST. The exam router includes both student and admin endpoints because the underlying concept is the exam session, and all related operations are collected there. While this creates a large file, it also means a reviewer can inspect most core product workflows in one place.",
                "The endpoint set reveals that the application supports warmup behavior, likely to reduce model initialization latency before an exam begins. That is an advanced usability feature because real-time computer-vision systems often feel unreliable when the first few seconds are spent loading detectors.",
            ],
        ),
        (
            "Authentication And Identity Model",
            [
                "Authentication is not limited to email and password. The backend supports password attempts, liveness challenges, password-plus-face confirmation, direct face login, LTI face binding, and initial password reset requirements. This makes the identity model significantly richer than a standard student portal.",
                "The presence of image-management endpoints under /me/image and face-enrollment analysis endpoints indicates that user identity is treated as a lifecycle. The student first becomes known to the system through enrollment media, then proves that identity through liveness and comparison at login, and may be rechecked during the exam itself.",
                "This layered identity design matters in proctoring. A single login event does not prove that the same person remains in front of the camera for the whole exam. The backend therefore performs recurring identity checks during frame processing, and detection_service keeps per-user streak state to control when a mismatch becomes a reportable alert.",
            ],
        ),
        (
            "Real-Time Proctoring Pipeline",
            [
                "detection_service.py is the clearest expression of the project’s technical ambition. It accepts frames, downsizes them for throughput control, evaluates frame quality, schedules detector execution, separates real-time detectors from deeper optional detectors, enforces a compute budget, and performs periodic identity verification against stored references.",
                "The detector strategy is intentionally adaptive. Face box, YOLO, and face mesh are treated as real-time lanes, while hand detection, gaze, and spoofing are opportunistic. This is smart because a proctoring system must remain responsive even when full inference would otherwise overload the host machine. The compute-budget guard shows that stability was treated as a first-class requirement.",
                "The service also carries anti-noise logic. It deduplicates logs, throttles detectors per user, suppresses brittle landmark-based checks when frame quality is poor, and introduces cooldowns and streak logic for identity alerts. These are the kinds of practical details that usually appear only after the team has experienced false positives and performance regressions.",
                "Evidence handling is integrated directly into processing flow. The service tracks upload retry windows and can degrade behavior when storage is unavailable. That design prevents a temporary object-store problem from fully crashing live monitoring, which is an important resilience characteristic.",
            ],
        ),
        (
            "Session Control, Logging, And Policy",
            [
                "The exam router and termination_policy_service together implement the governance model of the platform. Sessions can be started, paused, resumed, stopped, force-closed, summarized, and submitted. Violation categories are distinguished from system events, and the code defines major violation classes along with thresholds for automatic outcomes.",
                "The backend enforces exam status centrally. Helper methods fetch the latest or active session and explicitly deny access when the latest session has already been terminated. This is a strict but understandable policy: once the backend judges the attempt closed or invalid, the frontend cannot simply talk its way back into a valid state.",
                "log_service.py and models such as Log and Evidence provide the permanent record of what happened. This is crucial because a proctoring system is evaluated not only by what it detects in real time, but also by whether an instructor can later review, explain, and defend those detections.",
            ],
        ),
        (
            "Administrative And Reporting Capabilities",
            [
                "The backend includes a surprisingly broad admin data surface. There are endpoints for live sessions, room verification, exam-creator configuration, proctoring thresholds, result dashboards, student summaries, timeline views, evidence management, roster imports, and exam-link generation. This suggests that the project was designed as an end-to-end administrative tool rather than a narrow demo of computer vision.",
                "This admin surface is important for project evaluation because it converts machine-learning events into a usable operational workflow. Without these endpoints, suspicious detections would remain technical artifacts. With them, they become reviewable evidence, configurable policy, and actionable exam records.",
            ],
        ),
        (
            "Backend Engineering Takeaways",
            [
                "The backend demonstrates strong systems thinking. It balances hard real-time aspirations with practical throttling, uses services to keep responsibilities separated, and preserves exam truth on the server side. Its main maintainability challenge is breadth: exam.py is large, the project documentation has partial drift, and the amount of behavior encoded in backend state means testing discipline is especially important.",
                "A lecturer evaluating the backend should see it as the most technically dense part of the submission. It combines conventional web engineering with concurrency, media handling, rule systems, and machine-learning integration. Even where improvements are still possible, the architectural scope is substantial.",
            ],
        ),
    ]


def student_frontend_deep_dive() -> list[tuple[str, list[str]]]:
    return [
        (
            "Student Application Role",
            [
                "The student frontend is the exam client and candidate experience layer. It is not just a display shell around backend calls; it controls identity onboarding, pre-exam readiness, exam session initiation, live camera and screen-sharing orchestration, violation feedback, answer progression, and end-of-exam summarization.",
                "The React route map makes the intended journey very clear: login or signup, identity verification if needed, exam lobby, system check, network check, active exam, and summary. This route sequence mirrors the institutional workflow of online exams, which makes the product easier to explain to non-technical evaluators.",
            ],
        ),
        (
            "Authentication And Enrollment Experience",
            [
                "Login.jsx and Signup.jsx together form a guided identity workflow rather than a minimal sign-in screen. authService.js supports password login, login attempts, liveness challenge retrieval, password-plus-face login, LTI face binding, initial password reset, and face-enrollment analysis. The frontend therefore acts as a coordinator for several backend identity paths.",
                "Signup appears to emphasize guided reference capture from multiple face angles. This matters because later face comparison accuracy depends heavily on enrollment quality. From a project-understanding perspective, the signup flow is not peripheral; it is one of the foundations that makes later proctoring meaningful.",
                "VerifyIdentity.jsx extends that identity story into launch-time confirmation, especially for LTI or institutional flows. This allows the project to separate account existence from exam-session authenticity.",
            ],
        ),
        (
            "Lobby, Readiness, And Entry Control",
            [
                "ExamLobbyHome.jsx behaves like the student command center. It presents available exams, session readiness, and warmup-related behavior. In a proctoring product, the lobby is important because it is where the system can reduce failure before the exam timer starts.",
                "SystemCheckPage.tsx and NetworkCheckPage.tsx represent the bridge between conventional web navigation and secure exam mode. They check browser conditions, media-device readiness, and network/screen-sharing expectations. In the current branch, parts of the readiness flow appear simplified or bypassed to smooth demos and testing, which is worth understanding when comparing intended behavior with current behavior.",
                "The presence of lobbyProgress.js indicates that route access is not based only on the current URL. The frontend also tracks which prerequisites have been satisfied, which helps enforce stage order and creates a more guided student experience.",
            ],
        ),
        (
            "Active Exam Runtime",
            [
                "Exam.jsx is the operational heart of the student frontend. It coordinates exam questions, session timers, live WebSocket status, local media state, violation messaging, submission behavior, and admin-facing monitoring signals. Its size reflects the complexity of the live exam stage where many concerns must coexist in one screen.",
                "A major design detail lives in WebSocketHandler.js. The VideoStreamManager intentionally acquires the camera once and keeps it across socket reconnects so the student does not see camera flicker whenever network conditions change. That is an excellent user-experience choice because proctoring software often feels fragile when media devices repeatedly restart.",
                "Frame capture is adaptive. The handler tracks acknowledgements from the backend and adjusts capture cadence based on reported processing time. That means the student client is not blindly spamming frames; it cooperates with backend throughput. This is a sophisticated design touch and one of the most interesting frontend engineering decisions in the project.",
                "The hook layer, such as useScreenRecorder.js, shows that the frontend also treats screen sharing and recording as first-class behaviors. In other words, the active exam screen is not only a form or quiz interface; it is an orchestration hub for controlled monitoring.",
            ],
        ),
        (
            "Summary, Feedback, And Evidence Presentation",
            [
                "Summary.jsx closes the exam lifecycle by converting raw proctoring outcomes into a student-visible report. Compliance summaries, logged events, and export features such as PDF generation make the experience feel more complete and auditable.",
                "This final stage matters pedagogically as well. A good proctoring system should not only punish or flag; it should also explain what the system recorded. The summary page moves the platform closer to that ideal by making post-exam review a visible part of the user journey.",
            ],
        ),
        (
            "Frontend Engineering Takeaways",
            [
                "The student frontend demonstrates more than styling ability. It integrates routing, authentication state, media APIs, WebSocket coordination, answer/session services, and multi-stage exam readiness logic. Its main challenge is the same one faced by many real-world frontends: the active exam page accumulates a lot of responsibility and can become difficult to reason about over time.",
                "Still, the design intent is coherent. The app is built around the student journey, and the service layer provides a readable map of backend contracts. For an evaluator, this is evidence that the frontend was designed in relation to the product’s workflows rather than as isolated pages.",
            ],
        ),
    ]


def admin_frontend_deep_dive() -> list[tuple[str, list[str]]]:
    return [
        (
            "Admin Application Role",
            [
                "The admin frontend is the institutional control room of the platform. It exists because raw backend data and violation logs are not enough for practical proctoring. Instructors and administrators need live visibility, evidence review, result synthesis, threshold settings, and exam authoring tools.",
                "App.tsx shows a simple shell that pivots between admin login and the main dashboard, while AdminDashboard.tsx acts as the high-level navigation orchestrator. The design suggests that once authenticated, an admin stays inside one primary dashboard context and switches modules through tabs or sidebar-driven sections rather than separate standalone apps.",
            ],
        ),
        (
            "Live Monitoring And Exam Rooms",
            [
                "LiveMonitorLobby.tsx and LiveExamRoom.tsx are central to understanding the admin experience. The lobby appears to list active or available monitoring rooms, support room key verification, and transition into live session observation. This is the piece that turns the project from an automated checker into a supervised proctoring system.",
                "The naming and router support indicate that the admin side does not simply mirror the student webcam feed. Instead, it likely consumes a curated room abstraction, which is more scalable because admins typically monitor many sessions at once and only drill down when activity merits attention.",
            ],
        ),
        (
            "Results, Evidence, And Student Summaries",
            [
                "AdminResultsDashboard.tsx, EvidenceVault.tsx, and StudentSummaryPage.tsx represent the forensic and academic review side of the product. These modules convert logs, evidence files, session metrics, and incident timelines into a workflow an instructor can understand after the live exam has ended.",
                "This is an important strength in the project. Many proctoring demos stop at detecting suspicious activity, but an educational institution actually needs result interpretation and defensible evidence. The admin frontend acknowledges that requirement directly.",
            ],
        ),
        (
            "Exam Creation And Settings",
            [
                "ExamCreator.tsx makes the platform more self-contained because it lets admins define the assessments rather than assuming an external question bank will always exist. The backend endpoint inventory supports this with config endpoints, question creation, roster imports, and exam-link generation.",
                "AdminSettings.tsx pairs with backend settings and threshold endpoints. This suggests that the project treats policy and sensitivity values as operational controls rather than hard-coded constants hidden from administrators. That is a mature product decision because acceptable sensitivity may vary between institutions or exam types.",
            ],
        ),
        (
            "Design System And Composition",
            [
                "The admin app includes a large collection of UI components and utility wrappers, many inspired by Radix-style building blocks. This gives the application a composable structure where high-level pages are assembled from reusable controls, charts, dialogs, and layout helpers.",
                "A practical benefit of this componentized approach is that the admin side can grow without every new screen inventing its own interaction vocabulary. The tradeoff is that the file inventory becomes larger, which increases onboarding time. That is why the appendices in this report include a directory-by-directory catalog.",
            ],
        ),
        (
            "Admin Engineering Takeaways",
            [
                "The admin frontend rounds out the system and makes the whole project feel like a platform rather than a one-user demo. For a lecturer, it is strong evidence of full-stack scope because it shows that the team thought about how instructors investigate incidents, manage exams, and interpret results.",
                "Its main maintainability challenge is breadth rather than conceptual weakness. There are many modules, many UI helpers, and a lot of integration with backend contracts, so good documentation is essential. The built-in markdown reports found in the repo help compensate for that complexity and are included later in this document.",
            ],
        ),
    ]


def workflow_sections() -> list[tuple[str, list[str]]]:
    return [
        (
            "End-To-End Workflow: Student Signup To Exam Access",
            [
                "The likely onboarding story begins with student signup. The frontend collects conventional account data and face references, then calls backend enrollment analysis endpoints. The backend stores user records and reference imagery so later login and in-exam identity checks have trusted baseline material.",
                "At login, the student may pass through password-only, password-plus-face, or dedicated face-based flows depending on the policy branch. Liveness challenge endpoints indicate that the system can require the student to prove a live presence rather than submitting a static photo.",
                "Once authenticated, the student reaches the lobby where available exams are fetched, warmup may be triggered, and the frontend prepares for device and environment checks. This reduces the chance that model initialization or device access problems first appear after the exam has already started.",
            ],
        ),
        (
            "End-To-End Workflow: Exam Start And Active Monitoring",
            [
                "When the student enters the active exam stage, two parallel systems become important. First, the exam session exists as a backend record with start time, status, and related metadata. Second, the browser starts the real-time monitoring loop through camera acquisition, WebSocket connection, and frame capture.",
                "WebSocketHandler.js reveals a thoughtful connection strategy. The camera is acquired once and deliberately not reinitialized on socket reconnect. The client then sends JPEG frames and waits for backend acknowledgements. Processing time feedback influences future capture intervals, which helps synchronize client behavior with server capacity.",
                "On the backend, exam.py receives the session context while detection_service.py evaluates frames for face visibility, objects, gaze or mesh signals, spoofing, and identity mismatch. Logs are deduplicated and classified so that noise does not overwhelm the session record. When needed, evidence is persisted and serious incidents can influence policy outcomes.",
            ],
        ),
        (
            "End-To-End Workflow: Admin Observation And Post-Exam Review",
            [
                "While students are taking the exam, the admin frontend can inspect live rooms and session summaries using endpoints exposed by the backend. This means the system supports both automated policy and human oversight.",
                "After an exam, admin dashboards can retrieve results, timelines, evidence assets, and per-student summaries. This closes the loop between real-time monitoring and post-exam decision making. In institutional settings, this is the difference between a suspicious event being noticed and it actually becoming reviewable evidence.",
            ],
        ),
        (
            "Architecture Strengths, Risks, And Talking Points",
            [
                "Strong points to highlight to a lecturer include the multi-application scope, real-time proctoring architecture, multi-modal authentication, exam lifecycle completeness, evidence persistence, and operational admin tooling. These show breadth across frontend, backend, media APIs, storage, and ML integration.",
                "Honest risks to mention include documentation drift, large operational files such as exam.py and Exam.jsx, the challenge of reliably testing real-time detector behavior, and the need to keep current branch behavior aligned with intended exam-readiness policies. These are typical maturity issues for a project of this scale and do not erase its strengths.",
                "A good presentation framing is: this project is a full-stack AI proctoring platform that combines identity verification, exam management, real-time monitoring, and instructor review tools in one integrated system. That sentence is short, accurate, and easy for a lecturer to remember.",
            ],
        ),
    ]


def describe_endpoint(endpoint: dict[str, str]) -> str:
    path = endpoint["path"]
    handler = endpoint["handler"]
    method = endpoint["method"]
    router_file = endpoint["router_file"]

    if "auth" in router_file:
        base = "Authentication and identity workflow endpoint"
    elif "exam" in router_file:
        base = "Exam-session or admin exam management endpoint"
    elif "settings" in router_file:
        base = "System configuration or health endpoint"
    elif "observability" in router_file:
        base = "Observability and policy visibility endpoint"
    elif "lti" in router_file:
        base = "Learning Tools Interoperability integration endpoint"
    else:
        base = "Backend application endpoint"

    if "/admin/" in path:
        detail = "It is clearly oriented toward administrator workflows such as monitoring, results, settings, or exam authoring."
    elif "summary" in path or "timeline" in path or "evidence" in path:
        detail = "Its naming suggests review or evidence access rather than live transaction processing."
    elif any(token in path for token in ["/start", "/pause", "/resume", "/stop", "/submit"]):
        detail = "It directly participates in the exam lifecycle and likely mutates session state."
    else:
        detail = "It supports the normal application flow and likely coordinates several service and model calls."

    return f"{base}. Handler `{handler}` is exposed through {method} {path}. {detail}"


def describe_route(component: str, route_path: str) -> str:
    lower = component.lower()
    if "login" in lower:
        return "Login entry point for student authentication and transition into identity-aware access."
    if "signup" in lower:
        return "Enrollment route where a new student establishes account data and face references."
    if "verify" in lower:
        return "Identity verification step used after external launch or before protected exam access."
    if "lobby" in lower:
        return "Student dashboard and exam selection area where readiness is coordinated."
    if "systemcheck" in lower:
        return "Pre-exam device and browser validation page."
    if "networkcheck" in lower:
        return "Pre-exam connection and screen-sharing gate before entering the active exam."
    if component == "Exam":
        return "Primary exam-taking screen with live monitoring and answer workflow."
    if "summary" in lower:
        return "End-of-exam review page summarizing compliance and session results."
    return "Application route in the student exam journey."


def classify_file(root_name: str, path: Path) -> tuple[str, str]:
    rel = path.relative_to({"backend": BACKEND_ROOT, "student": STUDENT_ROOT, "admin": ADMIN_ROOT}[root_name])
    parts = rel.parts
    stem = path.stem
    text = read_text(path)

    category = "General Source"
    summary = f"{stem} contributes to the {root_name} application."

    if root_name == "backend":
        if parts[0] == "routers":
            category = "API Router"
            summary = f"{stem} defines a public HTTP contract area and maps inbound requests to backend workflows."
        elif parts[0] == "services":
            category = "Service Layer"
            summary = f"{stem} contains backend business logic that supports routes, policy, storage, or exam operations."
        elif parts[0] == "models":
            category = "ORM Model"
            classes = ", ".join(extract_python_classes(path)) or "database entities"
            summary = f"{stem} defines persistent SQLAlchemy entities such as {classes}, forming part of the exam data model."
        elif parts[0] == "schemas":
            category = "Pydantic Schema"
            classes = ", ".join(extract_python_classes(path)) or "request and response shapes"
            summary = f"{stem} defines API data contracts including {classes}, which shape validation and serialization."
        elif parts[0] == "detection":
            category = "Detection Module"
            summary = f"{stem} implements one part of the computer-vision detector stack used during live proctoring."
        elif parts[0] == "config":
            category = "Configuration"
            summary = f"{stem} centralizes runtime settings, database wiring, or detector tuning parameters."
        elif parts[0] == "utils":
            category = "Utility"
            summary = f"{stem} provides reusable backend helper behavior to keep routers and services focused."
        elif path.name == "main.py":
            category = "Application Entry Point"
            summary = "main.py constructs the FastAPI application, attaches middleware and routers, and exposes core startup behavior."

        special = {
            "main.py": "main.py is the backend bootstrap file. It initializes the FastAPI app, CORS, middleware, startup tasks, router registration, and WebSocket endpoints, so it is the best place to understand how the backend is assembled at runtime.",
            "exam.py": "exam.py is the largest workflow router in the backend. It covers warmup, student session operations, admin monitoring, evidence and summary access, exam creation, and several helper policies, making it the operational center of backend HTTP behavior.",
            "auth.py": "auth.py contains the identity and account API surface. It supports password login, face-based verification, enrollment analysis, profile image lifecycle, password reset requirements, and current-user access.",
            "detection_service.py": "detection_service.py is the real-time proctoring engine. It manages detector scheduling, frame-quality gating, compute budgets, identity rechecks, deduplication, and evidence upload coordination.",
            "termination_policy_service.py": "termination_policy_service.py encodes the escalation logic behind warnings, strike accumulation, and exam termination conditions. It turns raw events into institutional decisions.",
            "storage_service.py": "storage_service.py abstracts evidence persistence. It allows the backend to store files in MinIO-compatible storage while still supporting controlled fallback behaviors.",
            "log_service.py": "log_service.py converts live findings into durable audit entries. This is vital because a proctoring system is only trustworthy if incidents are reviewable after the live session ends.",
            "settings.py": "settings.py centralizes environment-driven application configuration such as URLs, secrets, storage settings, and enforcement options.",
            "detection_config.py": "detection_config.py concentrates threshold tuning for the proctoring models, making detector sensitivity easier to reason about and adjust.",
        }
        summary = special.get(path.name, summary)
    else:
        if "components" in parts:
            category = "React Component"
            summary = f"{stem} is a React component used to render UI and coordinate a slice of the {root_name} experience."
        elif "services" in parts:
            category = "Frontend Service"
            summary = f"{stem} wraps backend API calls or shared client-side workflow logic for the {root_name} app."
        elif "hooks" in parts:
            category = "Custom Hook"
            summary = f"{stem} packages reusable browser or stateful logic into a composable React hook."
        elif "store" in parts or "redux" in parts or "slices" in parts:
            category = "State Management"
            summary = f"{stem} helps manage shared application state across the {root_name} frontend."
        elif "utils" in parts:
            category = "Utility"
            summary = f"{stem} provides reusable client-side helpers to support UI logic and runtime behavior."
        elif "app" in parts:
            category = "App Module"
            summary = f"{stem} contributes to page structure or route-level behavior inside the {root_name} application."
        elif path.name == "App.jsx" or path.name == "App.tsx":
            category = "Application Shell"
            summary = f"{path.name} is the top-level routing shell for the {root_name} frontend."

        special = {
            "App.jsx": "App.jsx is the student application's routing shell. It maps the user journey from login and signup to identity verification, lobby, system checks, active exam, and final summary.",
            "Login.jsx": "Login.jsx coordinates the student sign-in experience, including password and face-aware authentication stages before protected exam access.",
            "Signup.jsx": "Signup.jsx guides a new student through registration and multi-angle face-reference capture so later identity verification has reliable enrollment material.",
            "VerifyIdentity.jsx": "VerifyIdentity.jsx is the launch-time identity confirmation screen, especially important for LTI or high-assurance exam entry flows.",
            "ExamLobbyHome.jsx": "ExamLobbyHome.jsx is the student dashboard for available exams and readiness progression. It is the natural bridge between authentication and the secure exam flow.",
            "SystemCheckPage.tsx": "SystemCheckPage.tsx verifies browser, device, and environment prerequisites before the student enters the monitored exam stage.",
            "NetworkCheckPage.tsx": "NetworkCheckPage.tsx covers the connection and screen-sharing gate immediately before the live exam. In the current branch, this flow appears simplified for smoother demo entry.",
            "Exam.jsx": "Exam.jsx is the most complex student screen. It coordinates question rendering, timing, WebSocket status, local media state, incident messaging, and submission behavior during the live exam.",
            "Summary.jsx": "Summary.jsx translates the completed exam session into a visible student-facing report with compliance context and export-oriented review.",
            "authService.js": "authService.js is the main client gateway for identity operations. It packages login, signup, liveness, face verification, and profile-related API calls for the student frontend.",
            "examService.js": "examService.js wraps session lifecycle, exam retrieval, progress persistence, logging, and submission APIs so the exam UI can stay focused on interaction rather than raw HTTP details.",
            "WebSocketHandler.js": "WebSocketHandler.js contains the VideoStreamManager used during live exams. It acquires the camera once, reconnects sockets without media flicker, adaptively captures frames, and coordinates with backend processing acknowledgements.",
            "useScreenRecorder.js": "useScreenRecorder.js manages screen-share and recording-oriented browser behavior needed for exam monitoring.",
            "App.tsx": "App.tsx is the admin application's top-level entry point. It gates the dashboard behind admin login and establishes the high-level structure of the admin experience.",
            "AdminDashboard.tsx": "AdminDashboard.tsx is the admin control shell that organizes monitoring, results, evidence, exam authoring, student summaries, and settings into one coherent workspace.",
            "LiveMonitorLobby.tsx": "LiveMonitorLobby.tsx is the admin entry point for active monitoring rooms. It translates live backend session data into a monitor-friendly workflow.",
            "LiveExamRoom.tsx": "LiveExamRoom.tsx is the drill-down live monitoring view where an administrator focuses on a specific exam room or student feed context.",
            "AdminResultsDashboard.tsx": "AdminResultsDashboard.tsx presents post-exam metrics, filtering, and analysis so suspicious sessions become reviewable academic results.",
            "EvidenceVault.tsx": "EvidenceVault.tsx organizes proctoring evidence retrieval and review, turning stored files and timelines into an investigator-friendly UI.",
            "ExamCreator.tsx": "ExamCreator.tsx lets administrators define exams and associated question data, helping make the platform self-contained.",
            "StudentSummaryPage.tsx": "StudentSummaryPage.tsx synthesizes a single candidate's exam and proctoring history into an admin-facing narrative view.",
            "AdminSettings.tsx": "AdminSettings.tsx provides the settings and threshold management interface that complements backend configuration endpoints.",
        }
        summary = special.get(path.name, summary)

    export_names = extract_export_names(path) if path.suffix.lower() in {".js", ".jsx", ".ts", ".tsx"} else []
    extra = ""
    if export_names:
        preview = ", ".join(export_names[:4])
        extra = f" Exported or detected symbols include {preview}."
    elif path.suffix.lower() == ".py":
        classes = extract_python_classes(path)
        if classes:
            extra = f" Declared classes include {', '.join(classes[:4])}."

    if "TODO" in text or "FIXME" in text:
        extra += " The file also contains explicit maintenance markers, suggesting active evolution."

    return category, (summary + extra).strip()


def add_manual_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title, 1)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)


def add_compound_sections(doc: Document, title: str, sections: list[tuple[str, list[str]]]) -> None:
    add_heading(doc, title, 1)
    for subtitle, paragraphs in sections:
        add_heading(doc, subtitle, 2)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)


def add_backend_catalog(doc: Document, endpoints: list[dict[str, str]]) -> None:
    add_heading(doc, "Appendix A: Backend Endpoint Catalog", 1)
    add_paragraph(
        doc,
        "This appendix lists the route surface detected in the backend router files. The descriptions are interpretive summaries generated from file grouping, method, and path naming, and they are meant to support understanding rather than replace full code review.",
    )
    rows: list[list[str]] = []
    for endpoint in endpoints:
        rows.append(
            [
                endpoint["router_file"],
                endpoint["method"],
                endpoint["path"],
                endpoint["handler"],
                describe_endpoint(endpoint),
            ]
        )
    add_table(doc, ["Router", "Method", "Path", "Handler", "Purpose"], rows)


def add_student_route_catalog(doc: Document, routes: list[dict[str, str]]) -> None:
    add_heading(doc, "Appendix B: Student Route Catalog", 1)
    rows = []
    for route in routes:
        rows.append(
            [
                route["path"],
                route["component"],
                describe_route(route["component"], route["path"]),
            ]
        )
    add_table(doc, ["Path", "Component", "Interpretation"], rows)


def add_model_catalog(doc: Document, model_files: list[Path]) -> None:
    add_heading(doc, "Appendix C: Backend Model Catalog", 1)
    add_paragraph(
        doc,
        "These entries summarize the persistence layer files under models/. In a project like this, the data model is important because exam truth, violations, evidence, and policy outcomes all become durable records.",
    )
    rows: list[list[str]] = []
    for path in model_files:
        classes = extract_python_classes(path)
        category, summary = classify_file("backend", path)
        rows.append(
            [
                str(path.relative_to(BACKEND_ROOT)),
                ", ".join(classes) or "No top-level class detected",
                category,
                summary,
            ]
        )
    add_table(doc, ["Model File", "Classes", "Type", "Commentary"], rows)


def add_file_inventory(doc: Document, title: str, root_name: str, files: list[Path]) -> None:
    add_heading(doc, title, 1)
    add_paragraph(
        doc,
        "This appendix catalogs source files so a reviewer can navigate the codebase systematically. Line counts are approximate and summaries are inference-based, but they provide a practical map for onboarding and oral presentation.",
    )

    grouped: dict[str, list[Path]] = defaultdict(list)
    for path in files:
        rel = path.relative_to({"backend": BACKEND_ROOT, "student": STUDENT_ROOT, "admin": ADMIN_ROOT}[root_name])
        group = rel.parts[0] if len(rel.parts) > 1 else "(root)"
        grouped[group].append(path)

    for group_name in sorted(grouped):
        add_heading(doc, f"{group_name} Directory", 2)
        rows: list[list[str]] = []
        for path in grouped[group_name]:
            category, summary = classify_file(root_name, path)
            rel_path = str(
                path.relative_to(
                    {"backend": BACKEND_ROOT, "student": STUDENT_ROOT, "admin": ADMIN_ROOT}[root_name]
                )
            )
            rows.append([rel_path, str(line_count(path)), category, summary])
        add_table(doc, ["File", "Lines", "Category", "Purpose Commentary"], rows)


def add_markdown_appendix(doc: Document, title: str, path: Path) -> None:
    add_heading(doc, title, 1)
    add_paragraph(
        doc,
        f"This appendix embeds and normalizes the repository document `{path.name}` so the final Word report preserves existing project documentation alongside the synthesized analysis.",
    )

    text = read_text(path)
    in_code = False
    code_buffer: list[str] = []

    def flush_code() -> None:
        nonlocal code_buffer
        if code_buffer:
            add_paragraph(doc, "\n".join(code_buffer))
            code_buffer = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)
            add_heading(doc, heading_match.group(2), level)
            continue

        if stripped.startswith(("- ", "* ")):
            add_paragraph(doc, stripped[2:], style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            continue

        add_paragraph(doc, stripped)

    flush_code()


def add_runtime_stack(doc: Document, backend_requirements: Path, student_package: Path, admin_package: Path) -> None:
    add_heading(doc, "Runtime Stack Snapshot", 1)
    add_paragraph(
        doc,
        "This section captures the current technical stack directly from the repository configuration files. It helps a reviewer explain the project in concrete engineering terms rather than only as a generic AI website.",
    )

    requirements_text = read_text(backend_requirements)
    package_student = load_json(student_package)
    package_admin = load_json(admin_package)

    backend_highlights = [
        "FastAPI, Uvicorn, and Gunicorn for the backend API and ASGI serving model.",
        "SQLAlchemy plus psycopg2-binary for relational persistence against PostgreSQL in the current compose setup.",
        "Redis for session or caching support and MinIO-compatible storage for evidence artifacts.",
        "OpenCV, MediaPipe, Ultralytics YOLO, TensorFlow CPU, Torch, DeepFace-related tooling, and other CV/ML libraries for proctoring intelligence.",
    ]

    student_dependencies = sorted((package_student.get("dependencies") or {}).keys())
    admin_dependencies = sorted((package_admin.get("dependencies") or {}).keys())

    add_heading(doc, "Backend Stack", 2)
    add_bullets(doc, backend_highlights)
    add_paragraph(
        doc,
        "requirements.txt contains both classic web dependencies and unusually heavy machine-learning dependencies, which reinforces that the backend is doing real proctoring work instead of only relaying messages between client and database.",
    )
    add_paragraph(doc, "The first lines of requirements.txt are reproduced below for context:")
    preview = "\n".join(requirements_text.splitlines()[:40])
    add_paragraph(doc, preview)

    add_heading(doc, "Student Frontend Stack", 2)
    add_paragraph(
        doc,
        "The student app is a React 18 + Vite application with Redux Toolkit, React Router, Chart.js, and a fairly rich set of Radix-based UI primitives. This suggests a modern client stack focused on interactive flows rather than static pages.",
    )
    add_paragraph(doc, "Representative dependencies detected in package.json:")
    add_bullets(doc, [", ".join(student_dependencies[:18]), ", ".join(student_dependencies[18:])])

    add_heading(doc, "Admin Frontend Stack", 2)
    add_paragraph(
        doc,
        "The admin app is also Vite-powered, but TypeScript-based and especially rich in dashboard-oriented and component-library dependencies, which matches its role as a configurable operations console.",
    )
    midpoint = max(1, len(admin_dependencies) // 2)
    add_bullets(doc, [", ".join(admin_dependencies[:midpoint]), ", ".join(admin_dependencies[midpoint:])])


def add_deployment_section(doc: Document, compose_path: Path, launcher_path: Path) -> None:
    compose_text = read_text(compose_path)
    launcher_text = read_text(launcher_path)

    add_heading(doc, "Deployment And Local Run Story", 1)
    add_paragraph(
        doc,
        "The repository supports both containerized backend execution and local frontend development servers. This mixed-mode developer experience is common in full-stack projects where the backend has heavy native dependencies while the frontends benefit from fast hot reload.",
    )
    add_paragraph(
        doc,
        "docker-compose.yml provisions PostgreSQL, Redis, the FastAPI backend, MinIO, and Adminer. The backend container exposes port 8080 while the admin and student frontends use Vite on ports 5173 and 5174 respectively. This makes the project easy to demo because each application has a stable, memorable port.",
    )
    add_paragraph(
        doc,
        "run_all.ps1 complements the container story by launching the three application tiers from a Windows PowerShell environment. Even if the script is mainly a convenience wrapper, it also documents the expected startup order and therefore helps explain the system architecture to a reviewer.",
    )

    add_heading(doc, "Compose Snapshot", 2)
    add_paragraph(doc, "\n".join(compose_text.splitlines()[:80]))
    add_heading(doc, "PowerShell Launcher Snapshot", 2)
    add_paragraph(doc, launcher_text)


def main() -> None:
    backend_files = list_source_files(BACKEND_ROOT)
    student_files = list_source_files(STUDENT_ROOT / "src")
    admin_files = list_source_files(ADMIN_ROOT / "src")

    backend_router_files = sorted((BACKEND_ROOT / "routers").glob("*.py"))
    model_files = sorted((BACKEND_ROOT / "models").glob("*.py"))
    endpoints: list[dict[str, str]] = []
    for router_file in backend_router_files:
        endpoints.extend(extract_backend_endpoints(router_file))

    student_routes = extract_react_routes(STUDENT_ROOT / "src" / "App.jsx")

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_manual_section(doc, "Executive Summary", app_kind_summary())
    add_manual_section(doc, "High-Level Architecture", architecture_summary())
    add_runtime_stack(
        doc,
        BACKEND_ROOT / "requirements.txt",
        STUDENT_ROOT / "package.json",
        ADMIN_ROOT / "package.json",
    )
    add_deployment_section(doc, ROOT / "docker-compose.yml", ROOT / "run_all.ps1")
    add_compound_sections(doc, "Backend Deep Dive", backend_deep_dive())
    add_compound_sections(doc, "Student Frontend Deep Dive", student_frontend_deep_dive())
    add_compound_sections(doc, "Admin Frontend Deep Dive", admin_frontend_deep_dive())
    add_compound_sections(doc, "Cross-System Workflows And Evaluation Notes", workflow_sections())

    add_backend_catalog(doc, endpoints)
    add_student_route_catalog(doc, student_routes)
    add_model_catalog(doc, model_files)
    add_file_inventory(doc, "Appendix D: Backend File Inventory", "backend", backend_files)
    add_file_inventory(doc, "Appendix E: Student Frontend File Inventory", "student", student_files)
    add_file_inventory(doc, "Appendix F: Admin Frontend File Inventory", "admin", admin_files)

    add_markdown_appendix(
        doc,
        "Appendix G: Embedded Admin Project Report",
        ADMIN_ROOT / "src" / "PROJECT_REPORT.md",
    )
    add_markdown_appendix(
        doc,
        "Appendix H: Embedded Admin Components Guide",
        ADMIN_ROOT / "src" / "ADMIN_COMPONENTS_GUIDE.md",
    )
    add_markdown_appendix(
        doc,
        "Appendix I: Embedded Backend README",
        BACKEND_ROOT / "README.md",
    )
    add_markdown_appendix(
        doc,
        "Appendix J: Embedded Student Frontend README",
        STUDENT_ROOT / "README.md",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Closing Note", 1)
    add_paragraph(
        doc,
        "This report is intended to help a lecturer understand the project as an integrated product: what each application does, how the runtime pieces interact, where the most important technical decisions live, and which files are worth opening first during evaluation or presentation.",
    )

    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
